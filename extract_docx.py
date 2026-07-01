#!/usr/bin/env python3
"""Extract case study text from romulus-v2.html and produce a clean .docx."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import html

def clean(text):
    """Strip HTML tags and decode entities."""
    text = re.sub(r'<[^>]+>', '', text)
    text = html.unescape(text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text, bold=False, italic=False):
    if not text.strip():
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

def add_bold_para(doc, label, rest=""):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    if rest:
        run2 = p.add_run(rest)
        run2.font.size = Pt(11)
    return p

def add_pullquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("· · ·")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Georgia'
font.size = Pt(11)

# ===== TITLE =====
doc.add_heading('Romulus', level=0)
p = doc.add_paragraph()
run = p.add_run('A Case Study')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
run = p.add_run('A case study in designing AI as infrastructure — and using it to ship four products in 87 days.')
run.italic = True
run.font.size = Pt(12)

p = doc.add_paragraph()
run = p.add_run('Feb 3 → May 1, 2026 · Brooklyn, NY · Mike Battaglia')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# ===== THE FRAME =====
add_heading(doc, 'The Frame', level=1)

add_para(doc, 'Most product teams treat AI as a tool. Open ChatGPT, paste a prompt, copy the result. The model is borrowed. The context resets every conversation. The compounding never starts.')

add_para(doc, 'I treated it as infrastructure — something with an identity, a memory architecture, a routing layer, and a delegation pattern. I designed it the way I\'d design any product: identity first, constraints defined, goals clear, then build. I run product at Daylight by day; this work happens nights and weekends. Designing the AI as infrastructure was the only way to ship four products in three months without quitting the job.')

add_para(doc, 'Week one, I rejected the trending pile. Boring problems, real money, speed to first dollar. Romulus had to support that bet, not just be impressive. From February 3 to May 1, I used what I built to ship a daily history game, a party-game thesis sharpened through multi-model cross-validation, a fullstack Gmail client (17,000 lines of code in two days, no keyboard time from me), a research-validated infrastructure thesis, and a multi-agent orchestration system that taught more than it shipped.')

add_pullquote(doc, 'The thesis is simple: designed AI compounds. Adopted AI doesn\'t. Everything that follows is the proof.')

# ===== ACT 1 — THE IDENTITY =====
doc.add_page_break()
add_heading(doc, 'Act 01 — The Identity', level=1)

add_para(doc, 'The name came first. Romulus — the founder-king of Rome. Not a conqueror. An architect of systems that outlasted him. The name wasn\'t decoration; it was the first design decision, and it set the tone for everything that followed.')

add_para(doc, 'Before I wrote a single prompt, I asked a harder question: what should this thing actually be?')

add_para(doc, 'The answer took the shape of SOUL.md — a product spec for a personality. Who is Romulus? What does he believe? How does he speak? What is he for? Then USER.md: my context, my background, my actual goals. Then IDENTITY.md, MEMORY.md, AGENTS.md, HEARTBEAT.md, TOOLS.md. All of it in place before a cron job ran, before a model was selected, before a line of code was written.')

add_para(doc, 'Inside SOUL.md, four contextual modes:')

add_bold_para(doc, 'Jarvis — ', 'Crisp reports. When speed matters.')
add_bold_para(doc, 'Consigliere — ', 'Strategic counsel. When stakes are high enough to push back.')
add_bold_para(doc, 'Cohort — ', 'Collaborative build mode. When a session is going sideways at midnight.')
add_bold_para(doc, 'Roman — ', 'Milestone gravitas. For the moments that matter.')

add_para(doc, '')
add_para(doc, 'And one hard rule, called The Fortress: nothing leaves the system without an explicit ask. No messages sent. No accounts created. No money spent. No production deploys. AI safety as a product design decision, not an afterthought.')

add_pullquote(doc, 'I designed AI the same way I\'d design any product. Identity first. Constraints defined. Goals clear. Then build.')

add_para(doc, 'The naming wasn\'t ornament — it scaled. Every cohort, every project, every artifact got a Roman name. Identity-gated, identity-named, identity-first. Ten years of building products for people taught me to start with the person using it. Here, the person was me. The product was my own operating system.')

# ===== ACT 2 — THE ARCHITECTURE =====
doc.add_page_break()
add_heading(doc, 'Act 02 — The Architecture', level=1)
add_para(doc, 'February 3, 2026. First boot. Three principles, earned upgrade by upgrade.', italic=True)

add_para(doc, 'MacBook Pro. OpenClaw installed. Discord wired to a private channel, #roma-eterna. Flat MEMORY.md as the only persistent memory. The first response landed clean. Then the system started breaking against itself, and the architecture had to be designed in response.')

# Principle 1
add_heading(doc, 'Principle 01 — Memory: vault as brain, sessions as work', level=2)

add_para(doc, 'Memory is the first design decision, not the last. MEMORY.md as a flat file was the equivalent of a notebook with no chapters. It worked for the first few days, then became unreadable. Every breakdown after that was a memory architecture failure in disguise.')

add_para(doc, 'The split that mattered emerged in four phases:')

items = [
    'Flat MEMORY.md. Sessions wrote here. The notebook with no chapters.',
    'Obsidian vault. Daily notes, project files, decisions, all wikilinked. A graph, not a list.',
    'QMD as the semantic layer over the vault. 340 files indexed. Context retrieved by meaning, not keyword.',
    'Wiki vault. 27 sources ingested. The shared memory layer the system queries before every session.',
]
for item in items:
    add_para(doc, item)

add_para(doc, '')
add_para(doc, 'Sessions are ephemeral. The vault persists. Every session writes to it; the next one reads from it before it opens. Context survives across weeks. The system stops resetting.')

add_pullquote(doc, 'The vault is the brain. Sessions are the work. That split is what makes the system compound over time rather than reset.')

add_para(doc, 'The same instinct shows up in trust boundaries. For weeks the morning brief delivered R train departures from Times Square instead of Prospect Ave — R16 looked right in the data; R34N is the actual Prospect Ave northbound stop. Plausible-looking data isn\'t verified data. In an AI system that synthesizes plausibility, the verification habit is the trust boundary. The lesson generalizes — product analytics, user research, AI output, all the same shape.')

add_divider(doc)

# Principle 2
add_heading(doc, 'Principle 02 — Routing: one model per job, costed', level=2)

add_para(doc, 'Model selection is infrastructure, not preference. I learned this on April 9 — a heavy Legion coding session on MiniMax M2.7. 205K context window, which seemed like enough. The session hit 338 messages and 2.5MB. Context overflowed. Three compaction events fired in sequence. Six consecutive edit-tool failures — the model was matching against stale text that no longer existed in the file. The session was gone.')

add_para(doc, 'The routing rules got written that night, in blood, into TOOLS.md and AGENTS.md. Each model gets a defined job, a cost budget, and a context ceiling.')

add_para(doc, '')
add_bold_para(doc, 'Qwen 3.6 Plus — ', 'Daily chat, crons, morning briefs · $0.40 / $2.00 per M')
add_bold_para(doc, 'Claude Code (Opus) — ', 'All coding. Subscription, never API. · flat')
add_bold_para(doc, 'Grok 4 Fast — ', 'Heavy sessions. 2M context window. · $0.20 / $0.50 per M')
add_bold_para(doc, 'MiniMax M2.7 — ', 'Demoted after April 9. Quick chat only.')
add_para(doc, '')

add_para(doc, 'Cost discipline followed. A balance-monitor cron runs three times a day. Target: under $15/day burn. Budget is a design constraint, the same shape as latency or context length.')

add_para(doc, 'The routing layer needed a machine that could hold it. March 30 — Mac Mini migration. A dedicated host, romulus-mini, always on. Auto-login, sleep disabled, FileVault on. The morning brief stopped disappearing when the laptop lid closed. Claude Code ran while I slept. Background work became actually background.')

add_divider(doc)

# Principle 3
add_heading(doc, 'Principle 03 — Delegation: the spec is the contract', level=2)

add_para(doc, 'The pattern that governs everything now:')

delegation_items = [
    'I give direction — the goal, the constraints, the judgment calls.',
    'Romulus writes the spec or prompt file.',
    'Claude Code implements against the codebase.',
    'Romulus verifies, commits, pushes through staging.',
    'Romulus reports back.',
]
for item in delegation_items:
    add_para(doc, item)

add_para(doc, '')
add_para(doc, 'I write specs and prompts. I don\'t write production code. A good spec produces a good build. A vague spec produces vague output. The spec is the contract between me and the implementer — same instinct as a PRD between PM and engineer.')

add_para(doc, 'Engineering rigor is part of the substrate. Every project goes through a staging branch first; Vercel auto-deploys staging on push; production only after review. Never push directly to main. Never run vercel --prod. Both rules learned from real incidents in February.')

add_para(doc, 'The architecture didn\'t arrive fully formed. It was earned, upgrade by upgrade, February to May — each change motivated by a real breakdown in the version before it. Identity, memory, routing, delegation. Four design decisions in stack order. The same order any product gets designed.')

# ===== ACT 3 — A TUESDAY =====
doc.add_page_break()
add_heading(doc, 'Act 03 — A Tuesday', level=1)

add_para(doc, 'The alarm doesn\'t fire at 8am. The morning brief does. NWS weather for Park Slope. R train departures from Prospect Ave northbound. Commute score. Three curated product signals, the rest of the noise filtered before it reaches me. I read it on the train.')

add_para(doc, 'At Daylight I run product — decisions, roadmaps, team, stakeholders. Demanding and full. Romulus is silent during the day; the system was designed for that.')

add_para(doc, 'By 10pm I\'m back on Discord. A product question. A half-formed spec. Responses come back in seconds and they carry context — not just from this conversation but from three weeks of decisions in the vault. A prompt file gets written. I type run it. By midnight there\'s a diff. By 2am a commit has landed. By morning a Vercel preview URL is waiting.')

add_para(doc, 'The vault from Tuesday shapes Wednesday\'s brief. The decision from this sprint informs the next one. Not because the software was updated — because the context accumulated. That\'s the compounding.')

# ===== RECEIPTS =====
add_heading(doc, 'Receipts', level=1)
add_para(doc, 'Real numbers from a real system. Each one tells a story.', italic=True)

add_para(doc, '')
add_bold_para(doc, '29 — ', 'Projects in vault · Wikilinked, daily-noted, decision-anchored')
add_bold_para(doc, '17K — ', 'LOC, 2 days · Via, fullstack — no keyboard time')
add_bold_para(doc, '338 — ', 'Messages, Apr 9 · The session that died, in receipts')
add_bold_para(doc, '$15 — ', 'Per-day burn · Production AI as a budget line')
add_para(doc, '')

# ===== ACT 4 — THE WORK =====
doc.add_page_break()
add_heading(doc, 'Act 04 — The Work', level=1)
add_para(doc, 'Five outputs. Each one anchored to a different capability the system enabled.', italic=True)

add_para(doc, 'Cherry Street Labs is the studio that houses all of it — the one-person software outfit at cherrystreetlabs.com. The site itself is its own argument: an ASCII fluid distortion hero in vanilla JS and Three.js, no build step, ~300 lines of code. Constraints as design inputs, not obstacles. The studio framing is deliberate — each project shares infrastructure, decisions, and lessons. What I learn on one makes the next sharper.')

# Project I — Chronicle
add_heading(doc, 'Project I — Chronicle', level=2)
add_para(doc, 'One day, on a phone, through Discord.', italic=True)
add_para(doc, 'Live at thischronicle.com', italic=True)

add_para(doc, 'One historical event per day. One image. Guess the year in three attempts. Monday is warm — two clues, guiding. Sunday has none. The difficulty arc isn\'t decoration, it\'s the retention mechanic. Designed in, not bolted on.')

add_para(doc, 'The spec was locked March 29, 2026. The build happened the next day — not at a desk. On my phone, through Discord, Romulus coordinating, Claude Code running in the background. By end of day: a working Next.js app deployed to Vercel, 90 puzzles seeded, complete scoring logic, share card, stats modal, full seven-day difficulty arc.')

add_para(doc, 'The longer play: Chronicle is built for a specific exit. localStorage-only state, the Wordle model pre-acquisition. Retention metrics instrumented for D7/D30/D90. No paid tier for an acquirer to unwind. NYT Games is a $9B business with no history game. That\'s the thesis.')

add_para(doc, 'Ship in a day, from a phone, while not at a desk. The delegation pattern works at speed.', bold=True)

# Project II — Forbidden
add_heading(doc, 'Project II — Forbidden', level=2)
add_para(doc, 'Research as product design.', italic=True)
add_para(doc, 'Live at playforbidden.com', italic=True)

add_para(doc, 'Forbidden started as a question about a gap in the market — the mobile party game space has Heads Up! and very little else, the tabletop space has Taboo, nobody has rethought it seriously for phones. There was clearly something there. The question was what, exactly.')

add_para(doc, 'I ran the initial product thesis through three AI models independently — Grok, ChatGPT, and Romulus — each asked to pressure-test it separately and come back with where it was wrong. The answers converged around a contradiction with the original premise. Forbidden isn\'t digital Taboo. It isn\'t a Heads Up clone. It\'s a constraint-based pass-the-phone party game: one player clues under verbal constraint while the group guesses simultaneously, in real time. The tension comes from steering a room to an answer while navigating the forbidden words.')

add_para(doc, 'The reframe happened April 8, before a single line of game logic was written. 862-card corpus across six themed decks. A fake-door funnel measures iOS intent on the web build. The web version is the signal layer; the iOS build waits for signal to justify it.')

add_pullquote(doc, 'The framing came from the research, not the starting assumption. The research changed the product shape before any code changed.')

add_para(doc, 'Pressure-test a thesis across multiple models before committing. Research-as-design at startup-killing speed.', bold=True)

# Project III — Via
add_heading(doc, 'Project III — Via', level=2)
add_para(doc, 'Superhuman, at a fraction of the price.', italic=True)
add_para(doc, 'Built April 17–18, 2026 · Awaiting deploy', italic=True)

add_para(doc, 'An AI-powered Gmail client. Fast inbox triage, read receipts, calendar integration. Liquid-glass design system matching the iOS 26 aesthetic. Edge runtime for pixel tracking. Smart triage running on gpt-4o-mini through OpenRouter.')

add_para(doc, 'Built in two days. 17,000 lines of code. 143 files. 18 API routes. 28 Vitest tests passing. Four phases plus three fix passes — P0, P1, P2. Six commits. None of it written by me; all of it directed by me.')

add_para(doc, 'Spec to working build through the delegation pattern, end to end. The spec was the contract. Claude Code did the implementation against it. Romulus verified, committed, pushed through staging. I reviewed the staging URL and made the calls. The pattern that worked on a one-day game scaled to a fullstack app with auth, OAuth, Prisma, and a real product surface.')

add_para(doc, 'The delegation pattern scales. From a Wordle-style game in a day to a fullstack app in two.', bold=True)

# Research IV — AgentForge
add_heading(doc, 'Research IV — AgentForge', level=2)
add_para(doc, 'Vercel for DeFi MCP servers.', italic=True)
add_para(doc, 'Researched April 28, 2026 · Build in queue', italic=True)

add_para(doc, 'Hosted MCP infrastructure for autonomous AI agents transacting onchain. The thesis: every DeFi MCP server today is a self-hosted proof of concept. Nobody owns the production backend slot — plug-and-play hosted MCP for agents. Meanwhile, 19% of all onchain transactions are now agent-driven. 76% of stablecoin transfer volume. 165M+ x402 transactions logged.')

add_para(doc, 'The research-to-decision pipeline ran in a single day. Three rounds of Grok validation. YC RFS Summer 2026 alignment. Competitor teardown across hosted vs self-hosted. X-signal analysis pulled the agent-volume numbers. Phase 1 scope locked: hosting + x402 billing + observability + security. 15 beta testers identified through X.')

add_para(doc, 'No code yet — that\'s the point. The case for building gets made before the building starts. Most of what kills early projects is the wrong thesis, not the wrong implementation. The system was designed to catch the thesis before it became code.')

add_para(doc, 'Pre-build research rigor at startup-killing speed. The thesis gets validated before the implementation begins.', bold=True)

# Postmortem V — Legion
add_heading(doc, 'Postmortem V — Legion', level=2)
add_para(doc, 'The architecture was sound. The substrate wasn\'t.', italic=True)
add_para(doc, 'Phase 2A in progress', italic=True)

add_para(doc, 'The most ambitious attempt: a multi-agent orchestration layer running in Discord. Five cohorts, each answering one question and handing off to the next. Caesar — is this worth building? Augustus — build and deploy it. Vespasian — map the path to $20K MRR. Trajan — distribute it. Romulus — the orchestrator, deciding what gets sequenced when. One Discord prompt, fifteen minutes later: a live deployed product, a wealth strategy, and a launch plan.')

add_para(doc, 'First successful run April 9. Caesar produced a complete competitive landscape on Forbidden — 72% confidence, 80% opportunity score, 12 competitors mapped — delivered to #legion. Caesar worked. Augustus worked on small builds. The Phase 1 architecture held.')

add_para(doc, 'What didn\'t hold: the seams. When the Noctis test — a native iOS Swift app — hit the fifteen-minute timeout on Augustus, the chain stopped. Vespasian and Trajan never ran. State didn\'t survive across handoffs. There was no retry logic, no timeout tiering scaled to build complexity, no eval layer to catch when a handoff had silently degraded.')

add_pullquote(doc, 'The architecture was sound. The execution substrate wasn\'t ready. Recognizing the difference is engineering maturity.')

add_para(doc, 'Phase 2A is the response: an effort-tier system (Tier 1 / Tier 2 / Tier 3 timeouts matched to build scope), durable task state, structured retry patterns, regression detection at the handoffs. None of it exotic. It\'s the careful, unsexy work that turns orchestration from impressive-in-the-right-conditions into reliable.')

add_para(doc, 'Knowing when the idea is right and the substrate isn\'t — and building the substrate next, not abandoning the idea.', bold=True)

# ===== ACT 5 — THE FORWARD VIEW =====
doc.add_page_break()
add_heading(doc, 'Act 05 — The Forward View', level=1)

add_para(doc, 'What\'s about to happen to small product teams: the operator who designs their own AI infrastructure becomes the operator who can run an entire product line. Not faster. Differently shaped.')

add_para(doc, 'This isn\'t a productivity story. It\'s about which work a single PM can hold. Roadmaps that update themselves against your decision history. Research synthesis run across three models in parallel before you commit. Overnight execution as the standing background. Institutional memory that used to require a team of five — and that grows with every project rather than resetting at every onboarding.')

add_para(doc, 'The next two years will sort early companies into two groups. The first group designs an AI infrastructure layer the way they\'d design any other system — identity, memory, routing, delegation, safety boundaries — and runs lean on the leverage that creates. The second group consumes tools off the shelf and ends up with committees adopting them. The first group will ship more, more cheaply, with smaller teams. The second will hire to keep up and won\'t catch up.')

add_pullquote(doc, 'I built Romulus to find out which group I wanted to be in. The answer became clear by week three.')

# ===== BYLINE =====
doc.add_paragraph()
add_divider(doc)
add_para(doc, 'Mike Battaglia', bold=True)
add_para(doc, 'Senior PM at Daylight · Founder, Cherry Street Labs')
add_para(doc, 'Park Slope, Brooklyn · mikebatts.net')
add_para(doc, 'Roma Eterna · MMXXVI', italic=True)

# Save
output_path = '/Users/michaelbattaglia/.openclaw/workspace/mikebatts/Romulus-Case-Study.docx'
doc.save(output_path)
print(f"Saved to: {output_path}")
